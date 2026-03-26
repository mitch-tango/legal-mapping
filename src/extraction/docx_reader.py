"""Word document text extraction with Track Changes awareness."""

from __future__ import annotations

import hashlib
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from lxml import etree


@dataclass
class DocxReadResult:
    """Result of reading a DOCX file."""
    text: str
    file_hash: str
    had_track_changes: bool
    page_count_estimate: int | None = None
    error: str | None = None


# Word XML namespace
_WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NSMAP = {"w": _WORD_NS}

# Heading style prefix
_HEADING_PREFIX = "Heading"

# Markdown heading markers by level
_HEADING_MARKERS = {1: "#", 2: "##", 3: "###", 4: "####", 5: "#####", 6: "######"}

# Approximate chars per page for rough page count estimate
_CHARS_PER_PAGE = 3000


def _compute_file_hash(file_path: str) -> str:
    """Compute SHA-256 hash of file contents."""
    h = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def _detect_track_changes(doc: Document) -> bool:
    """Check if the document contains Track Changes markup."""
    for para in doc.paragraphs:
        xml = para._element
        if xml.findall(f".//{{{_WORD_NS}}}del") or xml.findall(f".//{{{_WORD_NS}}}ins"):
            return True
    return False


def _resolve_track_changes(element: etree._Element) -> None:
    """Accept all track changes in-place on an XML element.

    - Strips <w:del> elements (deleted text) entirely
    - Unwraps <w:ins> elements (keeps inserted text content)
    """
    # Remove deletions
    for del_elem in element.findall(f".//{{{_WORD_NS}}}del"):
        parent = del_elem.getparent()
        if parent is not None:
            parent.remove(del_elem)

    # Unwrap insertions (keep content, remove wrapper)
    for ins_elem in element.findall(f".//{{{_WORD_NS}}}ins"):
        parent = ins_elem.getparent()
        if parent is not None:
            index = list(parent).index(ins_elem)
            for child in list(ins_elem):
                parent.insert(index, child)
                index += 1
            parent.remove(ins_elem)


def _extract_run_text(run) -> str:
    """Extract text from a run, marking bold and italic."""
    text = run.text or ""
    if not text:
        return ""
    if run.bold and run.italic:
        return f"***{text}***"
    if run.bold:
        return f"**{text}**"
    if run.italic:
        return f"*{text}*"
    return text


def _extract_paragraph_text(para) -> str:
    """Extract text from a paragraph with inline formatting."""
    parts = []
    for run in para.runs:
        parts.append(_extract_run_text(run))
    return "".join(parts)


def _get_heading_level(para) -> int | None:
    """Get heading level from paragraph style, or None if not a heading."""
    style_name = para.style.name if para.style else ""
    if style_name.startswith(_HEADING_PREFIX):
        try:
            level = int(style_name[len(_HEADING_PREFIX):].strip())
            return level
        except ValueError:
            return None
    return None


def _extract_table_text(table) -> str:
    """Convert a table to pipe-delimited text."""
    lines = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = cell.text.strip().replace("\n", " ")
            cells.append(cell_text)
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def read_docx(file_path: str) -> DocxReadResult:
    """Extract structured text from a Word document.

    Steps:
    1. Open with python-docx
    2. Detect and resolve Track Changes
    3. Extract text preserving structure (headings, lists, tables, formatting)
    4. Compute SHA-256 file hash

    Returns DocxReadResult with structured text or error.
    """
    path = Path(file_path)
    if not path.exists():
        return DocxReadResult(
            text="", file_hash="", had_track_changes=False,
            error=f"File not found: {file_path}",
        )

    try:
        file_hash = _compute_file_hash(file_path)
    except OSError as e:
        return DocxReadResult(
            text="", file_hash="", had_track_changes=False,
            error=f"Cannot read file: {e}",
        )

    try:
        doc = Document(file_path)
    except Exception as e:
        return DocxReadResult(
            text="", file_hash=file_hash, had_track_changes=False,
            error=f"Cannot parse DOCX: {e}",
        )

    # Detect and resolve Track Changes
    had_track_changes = _detect_track_changes(doc)
    if had_track_changes:
        for para in doc.paragraphs:
            _resolve_track_changes(para._element)

    # Build structured text output
    output_parts: list[str] = []

    # Interleave paragraphs and tables in document order
    # python-docx exposes body elements in order via doc.element.body
    body = doc.element.body
    for child in body:
        tag = etree.QName(child.tag).localname if isinstance(child.tag, str) else ""

        if tag == "p":
            # Find the matching paragraph object
            para = None
            for p in doc.paragraphs:
                if p._element is child:
                    para = p
                    break
            if para is None:
                continue

            heading_level = _get_heading_level(para)
            text = _extract_paragraph_text(para)

            if not text.strip():
                continue

            if heading_level and heading_level in _HEADING_MARKERS:
                output_parts.append(f"{_HEADING_MARKERS[heading_level]} {text}")
            else:
                output_parts.append(text)

        elif tag == "tbl":
            # Find matching table object
            for table in doc.tables:
                if table._element is child:
                    table_text = _extract_table_text(table)
                    if table_text.strip():
                        output_parts.append(table_text)
                    break

    full_text = "\n\n".join(output_parts)

    # Rough page count estimate
    page_estimate = max(1, len(full_text) // _CHARS_PER_PAGE) if full_text else 1

    return DocxReadResult(
        text=full_text,
        file_hash=file_hash,
        had_track_changes=had_track_changes,
        page_count_estimate=page_estimate,
    )
